import os
from pathlib import Path
from typing import List
import PyPDF2
from docx import Document as DocxDocument
from config import PROJECT_DOCS_PATH, FILE_EXTENSIONS, VECTOR_DB_PATH, CHUNK_SIZE, CHUNK_OVERLAP

class DocumentProcessor:
    def __init__(self):
        self.project_docs_path = PROJECT_DOCS_PATH
        self.supported_extensions = FILE_EXTENSIONS
    
    def load_all_documents(self) -> List[dict]:
        """Load all documents from project archive"""
        documents = []
        
        if not os.path.exists(self.project_docs_path):
            print(f"Warning: Project docs path does not exist: {self.project_docs_path}")
            return documents
        
        for root, dirs, files in os.walk(self.project_docs_path):
            for file in files:
                file_path = os.path.join(root, file)
                ext = Path(file_path).suffix.lower()
                
                if ext in self.supported_extensions:
                    print(f"Processing: {file_path}")
                    content = self._extract_text(file_path)
                    
                    if content:
                        documents.append({
                            'source': file_path,
                            'content': content,
                            'filename': file,
                            'extension': ext
                        })
        
        print(f"Loaded {len(documents)} documents")
        return documents
    
    def _extract_text(self, filepath: str) -> str:
        """Extract text from PDF, DOCX, or TXT files"""
        try:
            ext = Path(filepath).suffix.lower()
            
            if ext == '.pdf':
                return self._extract_pdf(filepath)
            elif ext == '.docx':
                return self._extract_docx(filepath)
            elif ext == '.txt':
                return self._extract_txt(filepath)
        
        except Exception as e:
            print(f"Error extracting text from {filepath}: {e}")
        
        return ""
    
    def _extract_pdf(self, filepath: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error reading PDF {filepath}: {e}")
        return text
    
    def _extract_docx(self, filepath: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = DocxDocument(filepath)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        except Exception as e:
            print(f"Error reading DOCX {filepath}: {e}")
        return text
    
    def _extract_txt(self, filepath: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading TXT {filepath}: {e}")
            return ""
    
    def chunk_documents(self, documents: List[dict]) -> List[dict]:
        """Split documents into chunks for embedding"""
        chunks = []
        
        for doc in documents:
            content = doc['content']
            source = doc['source']
            
            # Simple chunking strategy
            words = content.split()
            chunk_words = []
            
            for word in words:
                chunk_words.append(word)
                
                if len(chunk_words) >= CHUNK_SIZE:
                    chunk_text = ' '.join(chunk_words)
                    chunks.append({
                        'content': chunk_text,
                        'source': source,
                        'filename': doc['filename']
                    })
                    
                    # Keep overlap
                    chunk_words = chunk_words[CHUNK_SIZE - CHUNK_OVERLAP:]
            
            # Add remaining text
            if chunk_words:
                chunk_text = ' '.join(chunk_words)
                chunks.append({
                    'content': chunk_text,
                    'source': source,
                    'filename': doc['filename']
                })
        
        print(f"Created {len(chunks)} chunks from {len(documents)} documents")
        return chunks
